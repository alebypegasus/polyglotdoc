import logging
from pathlib import Path
from typing import List, Dict, Any, Callable, Optional
import docx
from backend_fastapi.app.services.ai_translator import ai_translator, DocumentTranslationContext

logger = logging.getLogger(__name__)

class DOCXProcessor:
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        # Approximate pages by paragraph count / word count
        total_paras = len(doc.paragraphs)
        pages = max(1, total_paras // 10)
        return {
            "total_pages": pages,
            "title": Path(file_path).stem,
            "author": "Document Author"
        }

    async def process_docx(
        self,
        input_path: str,
        output_path: str,
        target_language: str,
        source_language: str = "auto",
        progress_callback: Optional[Callable[[Dict[str, Any]], Any]] = None
    ) -> Dict[str, Any]:
        doc = docx.Document(input_path)
        context = DocumentTranslationContext(source_language=source_language, target_language=target_language)
        
        # Collect paragraphs
        blocks = []
        block_id = 0
        para_refs = []
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                blocks.append({"id": block_id, "text": text, "bbox": [0, 0, 100, 20]})
                para_refs.append(p)
                block_id += 1

        # Collect table cells
        cell_refs = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        blocks.append({"id": block_id, "text": text, "bbox": [0, 0, 100, 20]})
                        cell_refs.append(cell)
                        block_id += 1

        total_blocks = len(blocks)
        batch_size = 10
        chunk_count = max(1, (total_blocks + batch_size - 1) // batch_size)

        if progress_callback:
            await progress_callback({
                "status": "extracting",
                "current_page": 1,
                "total_pages": chunk_count,
                "percentage": 10.0,
                "message": "Extraindo parágrafos e tabelas do DOCX"
            })

        translated_blocks = []
        for i in range(0, total_blocks, batch_size):
            chunk = blocks[i : i + batch_size]
            current_page = (i // batch_size) + 1
            
            if progress_callback:
                pct = round(20 + (current_page / chunk_count) * 60, 1)
                await progress_callback({
                    "status": "translating",
                    "current_page": current_page,
                    "total_pages": chunk_count,
                    "percentage": pct,
                    "message": f"Traduzindo lote {current_page}/{chunk_count}"
                })

            trans_chunk = await ai_translator.translate_blocks(
                chunk, 
                context, 
                page_num=current_page, 
                total_pages=chunk_count
            )
            translated_blocks.extend(trans_chunk)

        if progress_callback:
            await progress_callback({
                "status": "reconstructing",
                "current_page": chunk_count,
                "total_pages": chunk_count,
                "percentage": 90.0,
                "message": "Reconstruindo arquivo DOCX final"
            })

        # Apply translations back to paragraphs
        trans_map = {b["id"]: b.get("translated_text", "") for b in translated_blocks}
        
        for idx, p in enumerate(para_refs):
            if idx in trans_map:
                # Keep initial run formatting if exists, replace text
                new_text = trans_map[idx]
                if p.runs:
                    p.runs[0].text = new_text
                    for r in p.runs[1:]:
                        r.text = ""
                else:
                    p.text = new_text

        # Apply to table cells
        cell_start_idx = len(para_refs)
        for idx, cell in enumerate(cell_refs):
            target_id = cell_start_idx + idx
            if target_id in trans_map:
                cell.text = trans_map[target_id]

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return {
            "total_pages": chunk_count,
            "output_path": output_path,
            "total_blocks": total_blocks
        }

docx_processor = DOCXProcessor()
